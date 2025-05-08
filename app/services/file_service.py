import os
from pathlib import Path
from typing import Optional, Union, List, Dict
import PyPDF2
import docx
from ..config import UPLOAD_DIR, ALLOWED_EXTENSIONS, MAX_FILE_SIZE
from ..models.document import Document
from app.models.structured_doc import SectionContent, DocumentStructure

class FileService:
    @staticmethod
    def validate_file(file) -> bool:
        """Validate file type and size."""
        if not file:
            return False
        
        file_extension = Path(file.name).suffix.lower()
        return (
            file_extension in ALLOWED_EXTENSIONS and
            file.size <= MAX_FILE_SIZE
        )

    @staticmethod
    def extract_text(file) -> Optional[Union[str, List[Union[str, Dict]]]]:
        """Extract text from different file types. Returns list for DOCX, string for others."""
        file_extension = Path(file.name).suffix.lower()
        
        try:
            if file_extension == '.pdf':
                return FileService._extract_from_pdf(file)
            elif file_extension == '.docx':
                return FileService._extract_from_docx(file)
            elif file_extension == '.txt':
                return file.getvalue().decode("utf-8")
        except Exception as e:
            print(f"Error extracting text: {e}")
            return None

    @staticmethod
    def _extract_from_pdf(file) -> str:
        text = ""
        pdf_reader = PyPDF2.PdfReader(file)
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text

    @staticmethod
    def _extract_from_docx(file) -> List[Union[str, Dict]]:
        doc = docx.Document(file)
        content_blocks = []
        # Track where we are in the doc (paragraphs and tables are interleaved)
        para_idx, table_idx = 0, 0
        total_paras = len(doc.paragraphs)
        total_tables = len(doc.tables)
        # Build a flat list of (type, obj) in document order
        elements = []
        for block in doc.element.body:
            if block.tag.endswith('}p'):
                elements.append(('p', doc.paragraphs[para_idx]))
                para_idx += 1
            elif block.tag.endswith('}tbl'):
                elements.append(('t', doc.tables[table_idx]))
                table_idx += 1
        for typ, obj in elements:
            if typ == 'p':
                text = obj.text.strip()
                if text:
                    content_blocks.append(text)
            elif typ == 't':
                # Extract table as list of lists and as text
                table_data = [[cell.text.strip() for cell in row.cells] for row in obj.rows]
                table_text = '\n'.join(['\t'.join(row) for row in table_data])
                content_blocks.append({
                    'type': 'table',
                    'text': table_text,
                    'data': table_data
                })
        return content_blocks 